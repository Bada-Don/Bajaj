import requests
import os
import tempfile
import magic
from PyPDF2 import PdfReader
from docx import Document
from typing import List
from io import BytesIO

class DocumentProcessor:
    def __init__(self, chunk_length: int = 600, chunk_overlap: int = 200):
        self.chunk_length = chunk_length
        self.chunk_overlap = chunk_overlap
    
    def extract_text_from_url(self, url: str) -> tuple:
        """Extract text from PDF or DOCX URL and return (text, temp_file_path)"""
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            try:
                response = requests.get(url, timeout=30)
                response.raise_for_status()
                temp_file.write(response.content)
                temp_file.flush()
                print(f"Debug: Temporary file created at {temp_file.name}")
                text = self._extract_text_from_file(temp_file.name)
                return text, temp_file.name
            except Exception as e:
                print(f"Debug: Error processing URL: {e}")
                raise
    
    def extract_text_from_buffer(self, file_buffer: 'BytesIO', filename: str) -> str:
        """Extract text from a file in memory"""
        # Get the file type based on content
        content_start = file_buffer.read(2048)
        file_buffer.seek(0)  # Reset position after reading
        
        mime = magic.Magic(mime=True)
        file_type = mime.from_buffer(content_start)
        
        text = ""
        if file_type == "application/pdf":
            reader = PdfReader(file_buffer)
            for page in reader.pages:
                text += page.extract_text() or ""
                
        elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                          "application/msword"]:
            doc = Document(file_buffer)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
                
        return text.strip()
    
    def _extract_text_from_file(self, file_path: str) -> str:
        """Internal method to extract text based on file type"""
        mime = magic.Magic(mime=True)
        file_type = mime.from_file(file_path)
        
        text = ""
        if file_type == "application/pdf":
            with open(file_path, "rb") as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() or ""
        
        elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                          "application/msword"]:
            with open(file_path, "rb") as f:
                doc = Document(f)
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        return text.strip()
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text from local file"""
        with open(file_path, 'rb') as f:
            return self.extract_text_from_buffer(BytesIO(f.read()), os.path.basename(file_path))
    
    def chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks"""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_length
            chunks.append(text[start:end])
            start += self.chunk_length - self.chunk_overlap
        
        return chunks